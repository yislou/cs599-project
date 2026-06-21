# -*- coding: utf-8 -*-
"""CS599 大作业报告生成器 - 从 content 文件导入并生成 Word 文档."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from scripts.report_content import COVER, TOC, CH1, CH2, CH3, CH4, CH5, CH6, CH7


# ─── Helper Functions ──────────────────────────────────

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h

def h2(doc, text):
    return doc.add_heading(text, level=2)

def para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.bold = bold
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(11)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
    pPr.append(shading)

def note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def make_table(doc, headers, data):
    table = doc.add_table(rows=len(data)+1, cols=len(headers), style='Light Grid Accent 1')
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    return table


# ─── Main Builder ──────────────────────────────────────

def build():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)

    # ============ COVER PAGE ============
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COVER['title'])
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COVER['project'])
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    doc.add_paragraph()

    fields = [
        ('课程名称', COVER['course']),
        ('项目名称', COVER['project']),
        ('方向', COVER['direction']),
        ('学号', COVER['student_id']),
        ('姓名', COVER['student_name']),
        ('专业', COVER['major']),
        ('指导教师', COVER['advisor']),
        ('提交日期', COVER['date']),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'{label}: {value}')
        r.font.size = Pt(14)

    doc.add_page_break()

    # ============ TOC ============
    h = doc.add_heading('目  录', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item, is_bold in TOC:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item)
        r.font.size = Pt(12)
        r.bold = is_bold

    doc.add_page_break()

    # ============ CH1 ============
    h1(doc, CH1['title'])
    h2(doc, CH1['1.1'])
    para(doc, CH1['p1'])
    para(doc, CH1['p2'])
    para(doc, CH1['p3'])

    h2(doc, CH1['1.2'])
    para(doc, CH1['p4'])
    para(doc, CH1['p5_title'], bold=True)
    for b in CH1['bullets']:
        bullet(doc, b)

    h2(doc, CH1['1.3'])
    para(doc, CH1['p6'])
    make_table(doc,
        ['层级', '技术选型', '选型理由'],
        [
            ['LLM', 'DeepSeek Chat API', '性价比高的国产大模型, OpenAI兼容接口'],
            ['Embedding', 'BAAI/bge-small-zh-v1.5', '本地运行, 中英双语, 512维, 无需API费用'],
            ['Agent框架', 'LangGraph', '显式状态图, ReAct模式, 可观测性优秀'],
            ['向量数据库', 'ChromaDB', '嵌入式设计, 零运维, 持久化存储'],
            ['UI', 'Streamlit', '最快构建聊天界面的Python框架'],
            ['文档解析', 'PyPDF + python-docx', '支持 PDF/TXT/MD/DOCX 多格式'],
            ['容器化', 'Docker', '一键部署, 环境隔离'],
        ]
    )
    note(doc, '表 1-1: 技术栈选型总览')

    doc.add_page_break()

    # ============ CH2 ============
    h1(doc, CH2['title'])
    para(doc, CH2['p1'])

    h2(doc, CH2['2.1'])
    para(doc, CH2['p2'])
    para(doc, CH2['p3'])

    h2(doc, CH2['2.2'])
    para(doc, CH2['p4'])
    para(doc, CH2['p5'])

    h2(doc, CH2['2.3'])
    para(doc, CH2['p6'])
    para(doc, CH2['p7'])

    doc.add_page_break()

    # ============ CH3 ============
    h1(doc, CH3['title'])
    h2(doc, CH3['3.1'])
    para(doc, CH3['p1'])

    arch = (
        "    +---------------------------------------------------+\n"
        "    |              Streamlit UI Layer                    |\n"
        "    |  +--------------+  +---------------------------+  |\n"
        "    |  |   Sidebar     |  |        Chat Area           |  |\n"
        "    |  |  - Upload     |  |  - Message History         |  |\n"
        "    |  |  - Status     |  |  - Agent Thinking (expand) |  |\n"
        "    |  |  - Manage     |  |  - Source Citations        |  |\n"
        "    |  +--------------+  +---------------------------+  |\n"
        "    +---------------------------------------------------+\n"
        "    |            Agent Layer (LangGraph)                 |\n"
        "    |  +----------+   +--------------+  +-----------+  |\n"
        "    |  |  Agent   |-->| Tool Router  |->| Tool Node |  |\n"
        "    |  |  Node    |<--|              |  | (Execute) |  |\n"
        "    |  +----------+   +--------------+  +-----------+  |\n"
        "    |      |  Conversation Memory (Sliding Window)      |\n"
        "    +---------------------------------------------------+\n"
        "    |                RAG Layer                           |\n"
        "    |  +----------+  +-----------+  +----------------+  |\n"
        "    |  |  Loader  |->| Embedder  |->| Vector Store   |  |\n"
        "    |  |(PyPDF,   |  |(BGE-zh)   |  | (ChromaDB)     |  |\n"
        "    |  | Docx2txt)|  |           |  |                |  |\n"
        "    |  +----------+  +-----------+  +----------------+  |\n"
        "    +---------------------------------------------------+\n"
        "    |            Infrastructure                          |\n"
        "    |  +----------+  +-----------+  +----------------+  |\n"
        "    |  |  Config  |  | DeepSeek  |  | Docker         |  |\n"
        "    |  | (.env)   |  | API       |  | Container      |  |\n"
        "    |  +----------+  +-----------+  +----------------+  |\n"
        "    +---------------------------------------------------+"
    )
    code(doc, arch)
    note(doc, '图 3-1: 系统四层架构总览')

    h2(doc, CH3['3.2'])
    para(doc, CH3['p2'])
    para(doc, CH3['p3'])

    react_flow = (
        "User: [提问]\n"
        "  |\n"
        "  v\n"
        "+-----------+\n"
        "| Agent     | Think: 需要检索相关文档\n"
        "| Node      | Action: search_knowledge_base(...)\n"
        "+-----+-----+\n"
        "      |\n"
        "      v\n"
        "+-----------+\n"
        "| Tool      | 执行语义检索 -> 返回 Top-4 相关片段\n"
        "| Node      |\n"
        "+-----+-----+\n"
        "      |\n"
        "      v\n"
        "+-----------+\n"
        "| Agent     | Think: 已获得足够上下文\n"
        "| Node      | -> 综合推理 -> 生成答案 + 引用来源\n"
        "+-----+-----+\n"
        "      |\n"
        "      v\n"
        "最终回答: 根据文档X, 核心结论是... (来源: report.pdf)\n"
        "      |\n"
        "      v\n"
        "存入 Conversation Memory (支持后续追问)"
    )
    code(doc, react_flow)
    note(doc, '图 3-2: Agent ReAct 交互流程')

    h2(doc, CH3['3.3'])
    para(doc, CH3['p4_title'], bold=True)
    para(doc, CH3['p5'])
    para(doc, CH3['p6'])

    make_table(doc,
        ['数据流阶段', '输入', '输出'],
        [
            ['文档加载', 'PDF/TXT/MD/DOCX 文件', 'Document 对象列表'],
            ['文档分块', 'Document 对象', '分块后 Document 对象 (~1000字/块)'],
            ['向量化', '文本字符串', '512维浮点向量'],
            ['语义检索', '查询文本', 'Top-K 相关 Document (含相似度分数)'],
        ]
    )
    note(doc, '表 3-1: 各数据流阶段的输入输出')

    doc.add_page_break()

    # ============ CH4 ============
    h1(doc, CH4['title'])

    h2(doc, CH4['4.1'])
    para(doc, CH4['p1'])
    agent_code = (
        "def create_agent() -> StateGraph:\n"
        "    llm = create_llm()\n"
        "    llm_with_tools = llm.bind_tools(ALL_TOOLS)\n"
        "\n"
        "    def agent_node(state):\n"
        "        response = llm_with_tools.invoke(state['messages'])\n"
        "        return {'messages': [response]}\n"
        "\n"
        "    def should_continue(state):\n"
        "        last = state['messages'][-1]\n"
        "        if last.tool_calls:\n"
        "            return 'tools'\n"
        "        return '__end__'\n"
        "\n"
        "    workflow = StateGraph(AgentState)\n"
        "    workflow.add_node('agent', agent_node)\n"
        "    workflow.add_node('tools', ToolNode(ALL_TOOLS))\n"
        "    workflow.set_entry_point('agent')\n"
        "    workflow.add_conditional_edges('agent', should_continue, {\n"
        "        'tools': 'tools', '__end__': END\n"
        "    })\n"
        "    workflow.add_edge('tools', 'agent')\n"
        "    return workflow.compile()"
    )
    code(doc, agent_code)
    note(doc, '代码 4-1: Agent 核心图构建 (src/agent/core.py)')

    h2(doc, CH4['4.2'])
    para(doc, CH4['p2'])
    tool_code = (
        "@tool\n"
        "def search_knowledge_base(query: str) -> str:\n"
        '    \"\"\"在知识库中语义搜索相关文档片段.\"\"\"\n'
        "    results = similarity_search(query, k=4)\n"
        "    for i, doc in enumerate(results, 1):\n"
        "        score = doc.metadata.get('score', 'N/A')\n"
        "        source = doc.metadata.get('source', 'unknown')\n"
        "        parts.append(f'[{i}] source: {source} | score: {score}')\n"
        "    return '\\n'.join(parts)"
    )
    code(doc, tool_code)
    note(doc, '代码 4-2: search_knowledge_base Tool (src/agent/tools.py)')
    para(doc, CH4['p3'])

    h2(doc, CH4['4.3'])
    para(doc, CH4['p4'])
    config_code = (
        "# .env 文件结构 (实际值不提交到仓库)\n"
        "DEEPSEEK_API_KEY=sk-xxxxxxxx\n"
        "LLM_MODEL=deepseek-chat\n"
        "EMBEDDING_MODEL=BAAI/bge-small-zh-v1.5\n"
        "CHUNK_SIZE=1000\n"
        "MAX_HISTORY_TURNS=10"
    )
    code(doc, config_code)
    note(doc, '代码 4-3: 环境变量配置示例 (.env.example)')
    para(doc, CH4['p5'])

    doc.add_page_break()

    # ============ CH5 ============
    h1(doc, CH5['title'])

    h2(doc, CH5['5.1'])
    para(doc, CH5['p1'])
    make_table(doc,
        ['测试用例', '测试内容', '结果'],
        [
            ['Config Loading', '验证 .env 加载、路径解析、默认值', '通过'],
            ['Document Chunking', '验证文档分块数量、元数据保留', '通过'],
            ['Vector Store CRUD', '验证添加、检索、列表、删除、清空', '通过'],
            ['Agent Compilation', '验证 LangGraph 状态图编译成功', '通过'],
        ]
    )
    note(doc, '表 5-1: 集成测试结果 (4/4 通过)')

    h2(doc, CH5['5.2'])
    para(doc, CH5['p2'])
    para(doc, CH5['p3'])

    h2(doc, CH5['5.3'])
    para(doc, CH5['p4'])
    para(doc, CH5['p5'])

    doc.add_page_break()

    # ============ CH6 ============
    h1(doc, CH6['title'])

    h2(doc, CH6['6.1'])
    para(doc, CH6['p1'])

    h2(doc, CH6['6.2'])
    para(doc, CH6['p2_title'], bold=True)
    make_table(doc,
        ['优先级', '功能', '技术方案'],
        [
            ['P0', 'MCP 协议集成', '通过 MCP Server 外挂工具, Agent 可访问外部API/数据库'],
            ['P0', 'REST API', 'FastAPI 提供 HTTP 接口, 支持第三方系统集成'],
            ['P1', '用户认证', 'JWT + 多用户隔离的知识库 Collection'],
            ['P1', '云部署', 'Docker + 云服务器, 提供公网可访问 URL'],
            ['P2', '评估体系', 'RAGAS 评估框架 + 检索质量 Benchmark'],
        ]
    )
    note(doc, '表 6-1: v0.2 版本迭代计划')

    h2(doc, CH6['6.3'])
    para(doc, CH6['p3'])
    para(doc, CH6['p4'])

    doc.add_page_break()

    # ============ CH7 ============
    h1(doc, CH7['title'])

    h2(doc, CH7['7.1'])
    para(doc, CH7['p1'])
    para(doc, CH7['p2'])

    h2(doc, CH7['7.2'])
    para(doc, CH7['p3'])
    para(doc, CH7['p4'])

    h2(doc, CH7['7.3'])
    para(doc, CH7['p5'])
    para(doc, CH7['p6'])

    # ─── Save ─────────────────────────────────────────
    import os
    os.makedirs('docs', exist_ok=True)
    out = 'docs/CS599_大作业报告.docx'
    doc.save(out)
    print(f'OK -> {out}')


if __name__ == '__main__':
    build()
